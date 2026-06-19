import streamlit as st
from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import io

# Import untuk kebutuhan PDF generator secara native di server cloud
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Mengatur margin di dalam cell tabel menjadi 0 dxa agar pas dengan ukuran"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_table_borders(table):
    """Menghilangkan seluruh garis pembatas tabel di Word"""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# --- Konfigurasi Halaman Streamlit ---
st.set_page_config(page_title="Sticker Layout Generator", page_icon="🖨️")
st.title("🖨️ Generator Layout Stiker A4 Pro")
st.write("Layout stiker otomatis penuh satu lembar A4 dengan jeda horizontal antar stiker sebesar 1cm.")

# --- Input User ---
# Evaluasi 2: Menerima PNG, JPG, dan JPEG
uploaded_file = st.file_uploader("Upload Gambar Stiker (PNG / JPG / JPEG)", type=["png", "jpg", "jpeg"])
sticker_size_cm = st.number_input("Ukuran Sisi Stiker Persegi (cm)", min_value=1.0, max_value=20.0, value=5.0, step=0.1)

if uploaded_file is not None:
    image = Image.open(uploaded_file)
    st.image(image, caption="Preview Gambar Stiker", use_container_width=False, width=200)
    
    # --- Perhitungan Matematika Layout dengan Jeda 1cm ---
    printable_width = 21.0 - 2.0  # Total lebar area cetak (Margin kiri-kanan masing-masing 1cm)
    printable_height = 29.7 - 2.0 # Total tinggi area cetak (Margin atas-bawah masing-masing 1cm)
    gap_cm = 1.0                  # Jeda horizontal antar gambar
    
    # Menghitung jumlah kolom stiker yang muat dengan rumus: 
    # (Lebar cetak + gap) // (Ukuran stiker + gap)
    cols_count = int((printable_width + gap_cm) // (sticker_size_cm + gap_cm))
    rows_count = int(printable_height // sticker_size_cm)
    
    # Hitung total kolom aktual pada struktur tabel Word (termasuk kolom spacer kosong)
    total_word_cols = (cols_count * 2) - 1 if cols_count > 0 else 0
    
    if cols_count == 0 or rows_count == 0:
        st.error("Ukuran stiker atau jeda terlalu besar untuk area kertas A4!")
    else:
        st.info(f"📊 **Analisis Layout:** Terdeteksi {rows_count} Baris × {cols_count} Kolom Stiker (Total: {rows_count * cols_count} Stiker per Lembar).")
        
        # Buat pilihan format output
        # Evaluasi 3: Pilihan format Word atau PDF
        output_format = st.radio("Pilih Format Dokumen Output:", ("Microsoft Word (.docx)", "Portable Document Format (.pdf)"))
        
        if st.button("Proses dan Ambil File"):
            with st.spinner("Sedang menyusun layout stiker..."):
                
                # Standarisasi gambar ke byte stream agar aman dibaca docx/reportlab
                img_byte_arr = io.BytesIO()
                img_format = image.format if image.format else 'PNG'
                image.save(img_byte_arr, format=img_format)
                img_byte_arr.seek(0)
                
                # ----------------------------------------------------
                # JALUR GENERATE WORD DENGAN JEDA HORIZONTAL
                # ----------------------------------------------------
                if output_format == "Microsoft Word (.docx)":
                    doc = Document()
                    
                    # Atur spesifikasi kertas A4 dan margin 1cm
                    for section in doc.sections:
                        section.page_width = Cm(21.0)
                        section.page_height = Cm(29.7)
                        section.top_margin = Cm(1.0)
                        section.bottom_margin = Cm(1.0)
                        section.left_margin = Cm(1.0)
                        section.right_margin = Cm(1.0)
                    
                    # Membuat tabel dengan kolom spacer terintegrasi
                    table = doc.add_table(rows=rows_count, cols=total_word_cols)
                    remove_table_borders(table)
                    
                    for r in range(rows_count):
                        table.rows[r].height = Cm(sticker_size_cm)
                        
                        stiker_idx = 0
                        for c in range(total_word_cols):
                            cell = table.cell(r, c)
                            set_cell_margins(cell)
                            
                            if c % 2 == 0:
                                # Kolom Genap: Tempat menaruh stiker
                                cell.width = Cm(sticker_size_cm)
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(img_byte_arr, width=Cm(sticker_size_cm), height=Cm(sticker_size_cm))
                            else:
                                # Kolom Ganjil: Sebagai pembatas horizontal (Spacer 1cm)
                                cell.width = Cm(gap_cm)
                                
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    st.success("File Word berhasil dibuat!")
                    st.download_button(
                        label="📥 Download File Word (.docx)",
                        data=doc_buffer,
                        file_name="layout_stiker.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # ----------------------------------------------------
                # JALUR GENERATE PDF NATIVE
                # ----------------------------------------------------
                else:
                    pdf_buffer = io.BytesIO()
                    # Setup dokumen dengan margin 1cm (1cm = ~28.34 points)
                    margin = 1 * cm
                    doc_pdf = SimpleDocTemplate(
                        pdf_buffer, 
                        pagesize=A4,
                        leftMargin=margin, 
                        rightMargin=margin, 
                        topMargin=margin, 
                        bottomMargin=margin
                    )
                    
                    # Menyiapkan konfigurasi lebar kolom untuk tabel PDF
                    col_widths = []
                    for c in range(total_word_cols):
                        if c % 2 == 0:
                            col_widths.append(sticker_size_cm * cm)
                        else:
                            col_widths.append(gap_cm * cm)
                    
                    # Memasukkan element gambar ke dalam struktur matriks tabel PDF
                    data_matrix = []
                    for r in range(rows_count):
                        row_data = []
                        for c in range(total_word_cols):
                            if c % 2 == 0:
                                # Gunakan internal path/stream dari asset gambar
                                img_flowable = RLImage(img_byte_arr, width=sticker_size_cm*cm, height=sticker_size_cm*cm)
                                row_data.append(img_flowable)
                            else:
                                row_data.append("") # Cell kosong untuk jeda horizontal
                        data_matrix.append(row_data)
                    
                    # Styling tabel PDF agar margin internal cell bernilai 0 dan tanpa garis
                    pdf_table = Table(data_matrix, colWidths=col_widths, rowHeights=[sticker_size_cm*cm]*rows_count)
                    pdf_table.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ('LEFTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (0,0), (-1,-1), 0),
                        ('TOPPADDING', (0,0), (-1,-1), 0),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ]))
                    
                    # Bangun PDF ke dalam buffer memori RAM
                    story = [pdf_table]
                    doc_pdf.build(story)
                    pdf_buffer.seek(0)
                    
                    st.success("File PDF berhasil dibuat!")
                    st.download_button(
                        label="📥 Download File PDF (.pdf)",
                        data=pdf_buffer,
                        file_name="layout_stiker.pdf",
                        mime="application/pdf"
                    )
